import random
from flask import Blueprint, render_template, redirect, request, jsonify, session
from transformers import AutoTokenizer, AutoModel
import torch
from scipy.spatial.distance import cosine
import spacy
from llama_cpp import Llama
from docx import Document

llm = Llama(model_path="tinyllama-1.1b-chat-v0.3.Q5_K_M.gguf")#Please replace to your LLM file path

COLOR_MODEL = spacy.load("COLOR_model")
OBJECT_MODEL=spacy.load("OBJECT_model")
FORM_MODEL=spacy.load("FORM_model")
BOLD_MODEL=spacy.load("BOLD_model")
TOPIC_MODEL=spacy.load("TOPIC_model")
app = Blueprint("agent", __name__)


# Preparing the model and tokenizer
model_name = "bert-base-multilingual-cased"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModel.from_pretrained(model_name)

data_dic = {
"color": "",
"change_object1": "",
"change_object2": "",
"play_animation":"",
"play_line_animation":"",
"form_text":"",
"new_object":"",
"new_object_color":"",
"delete_object":"",
"bold_text":"",
"word_file":"",
}  

def reset_all():
    for key in data_dic:
        data_dic[key] = ""

def compute_similarity(task1, task2):
    # Tokenize the task and convert it into a format that can be input into the model
    inputs1 = tokenizer(task1, return_tensors='pt', truncation=True, padding=True)
    inputs2 = tokenizer(task2, return_tensors='pt', truncation=True, padding=True)
    # Compute embeddings for each task
    with torch.no_grad():
        outputs1 = model(**inputs1)
        embeddings1 = torch.mean(outputs1.last_hidden_state, dim=1).squeeze()
        outputs2 = model(**inputs2)
        embeddings2 = torch.mean(outputs2.last_hidden_state, dim=1).squeeze()
    # Calculate similarity (obtain similarity by subtracting cosine distance from 1)
    similarity = 1 - cosine(embeddings1.cpu(), embeddings2.cpu())
    return similarity


def object_similarity(object_a, object_b):
    # Tokenize a string and convert it into a format that can be input into a model
    inputs1 = tokenizer(object_a, return_tensors='pt', truncation=True, padding=True)
    inputs2 = tokenizer(object_b, return_tensors='pt', truncation=True, padding=True)
    # Extract features of each string using a model
    with torch.no_grad():
        outputs1 = model(**inputs1)
        outputs2 = model(**inputs2)
    # Get the features of the first token ([CLS])
    embeddings1 = outputs1.last_hidden_state[:, 0, :]
    embeddings2 = outputs2.last_hidden_state[:, 0, :]
    # Remove extra dimensions if embeddings are not 1-dimensional
    if len(embeddings1.shape) > 1:
        embeddings1 = embeddings1.squeeze()
    if len(embeddings2.shape) > 1:
        embeddings2 = embeddings2.squeeze()
    # Calculate cosine similarity
    similarity = 1 - cosine(embeddings1, embeddings2)
    #print(f"Similarity between {object_a} and {object_b}: {similarity}")
    return similarity


# Defining tasks and their corresponding functions
def task1(most_similar_task):
    reset_all()
    object_dict = {
        "blue button": "blue-button",
        "gray button": "gray-button",
        "green button": "green-button",
        "light blue button": "light-blue-button",
        "red button": "red-button",
        "yellow button": "yellow-button",
        "white button": "white-button",
        "black button": "black-button",
        "input form": "text-form",
        "text form": "text-form",
    }
    response_list = ["I changed element color.",
                     "I modified the element color.",
                     "I transform the element color."]
    response = random.choice(response_list)
    max_similarity = -1
    max_similarity_key = None
    #Get key
    left_values = list(object_dict.keys())
    #get value
    right_values = list(object_dict.values())
    #Loading the model
    doc_object = OBJECT_MODEL(most_similar_task)
    object = next((ent_obj.text for ent_obj in doc_object.ents if ent_obj.label_ == "OBJECT"), None)
    updated_sentence = most_similar_task.replace(object, "")
    for l in left_values:
        similarity = object_similarity(object, l)
        if similarity > max_similarity:
            max_similarity = similarity
            max_similarity_key = right_values[left_values.index(l)]
            print("1",max_similarity_key)
        if max_similarity > 0.9:
            print("1_2",max_similarity_key)
            data_dic["change_object1"] = max_similarity_key
    print(updated_sentence)
    doc_color = COLOR_MODEL(updated_sentence)
    color = next((ent.text for ent in doc_color.ents if ent.label_ == "COLOR"), None)
    print("Detected colors:", color)
    data_dic["color"] = color
    print("data_dic", data_dic)
    return response, data_dic

#Enter text into form
def task2(most_similar_task):
    reset_all()
    response_list = ["I entered a value in the form.",
                     "I inputted a value into the form.",
                     "I filled in a value on the form."]
    response = random.choice(response_list)
    doc = FORM_MODEL(most_similar_task)
    print(f"Text: {most_similar_task}")
    updated_task = most_similar_task
    for ent in doc.ents:
        print(f"{ent.text} ({ent.label_})")
        updated_task = updated_task.replace(ent.text, '')
    updated_task = updated_task.replace(".", '')
    print(f"Updated Text: {updated_task}")
    data_dic["form_text"] = updated_task
    return response, data_dic

def task3(most_similar_task):
    reset_all()
    object_dict = {
        "blue button": "blue-button",
        "gray button": "gray-button",
        "green button": "green-button",
        "light blue button": "light-blue-button",
        "red button": "red-button",
        "yellow button": "yellow-button",
        "white button": "white-button",
        "black button": "black-button",
        "input form": "text-form",
        "text form": "text-form",
    }
    response_list = ["I changed the arrangement of the elements.",
                     "I altered the order of the elements.",
                     "I adjusted how the elements were organized."]
    response = random.choice(response_list)
    max_similarity = -1
    max_similarity_key = None
    #Get key
    left_values = list(object_dict.keys())
    #Get value
    right_values = list(object_dict.values())
    doc = OBJECT_MODEL(most_similar_task)
    for i, ent in enumerate(doc.ents):
        print(f"{ent.text} ({ent.label_})")
        if i == 0:
            # Calculate the similarity with all elements and select the one with the highest similarity
            for l in left_values:
                similarity = object_similarity(ent.text, l)
                if similarity > max_similarity:
                    max_similarity = similarity
                    max_similarity_key = right_values[left_values.index(l)]
                    print("1",max_similarity_key)
            if max_similarity > 0.9:
                print("1_2",max_similarity_key)
                data_dic["change_object1"] = max_similarity_key
        elif i == 1:
            #Initializing values for data extraction
            max_similarity = -1
            max_similarity_key = None
            for l in left_values:
                similarity = object_similarity(ent.text, l)
                if similarity > max_similarity:
                    max_similarity = similarity
                    max_similarity_key = right_values[left_values.index(l)]
                    print("2",max_similarity_key)
            if max_similarity > 0.9:
                print("2_2",max_similarity_key)
                data_dic["change_object2"] = max_similarity_key
    print(data_dic)
    return response, data_dic


def task4(most_similar_task):
    reset_all()
    response_list = ["I created a new element.",
                     "I generated a fresh element.",
                     "I designed a newly-introduced element."]
    response = random.choice(response_list)
    if "button" in most_similar_task:
        color= None
        doc_color = COLOR_MODEL(most_similar_task)
        color = next((ent.text for ent in doc_color.ents if ent.label_ == "COLOR"), None)
        data_dic["new_object"] = "button"
        if color:
            data_dic["new_object_color"] = color
    elif "form" in most_similar_task:
        data_dic["new_object"] = "text-form"
    return response, data_dic

def task5(most_similar_task):
    reset_all()
    object_dict = {
    "blue button": "blue-button",
    "gray button": "gray-button",
    "green button": "green-button",
    "light blue button": "light-blue-button",
    "red button": "red-button",
    "yellow button": "yellow-button",
    "white button": "white-button",
    "black button": "black-button",
    "input form": "text-form",
    "text form": "text-form",
    "new button": "new_button",
    "new button2": "new_button2",
    "new button3": "new_button3",
    "new button4": "new_button4",
    }
    response_list = ["OK, I deleted an element.",
                     "I removed an element.",
                     "I took out an element."]
    response = random.choice(response_list)
    max_similarity = -1
    max_similarity_key = None
    left_values = list(object_dict.keys())
    right_values = list(object_dict.values())
    doc = OBJECT_MODEL(most_similar_task)
    #Allows you to retrieve multiple elements
    for i, ent in enumerate(doc.ents):
        print(f"{ent.text} ({ent.label_})")
        if i == 0:
            # Calculate the similarity with all elements and select the one with the highest similarity
            for l in left_values:
                similarity = object_similarity(ent.text, l)
                if similarity > max_similarity:
                    max_similarity = similarity
                    max_similarity_key = right_values[left_values.index(l)]
                    print("1",max_similarity_key)
            if max_similarity > 0.9:
                print("1_2",max_similarity_key)
                data_dic["delete_object"] = max_similarity_key
    print(data_dic)
    return response, data_dic

def task6(most_similar_task):
    reset_all()
    response_list = ["OK, let's play line animation.",
                     "Alright, let's initiate the line animation.",
                     "Okay, it's time to start the line animation."]
    response = random.choice(response_list)
    data_dic["play_line_animation"] = "play"
    print(data_dic)
    return response, data_dic

def task6_2(most_similar_task):
    reset_all()
    response_list = ["I stoped animation.",
                     "I ended the animation.",
                     "I paused the animation."]
    response = random.choice(response_list)
    data_dic["play_line_animation"] = "stop"
    print(data_dic)
    return response, data_dic
 
def task7(most_similar_task):
    reset_all()
    response_list = ["Sure, let's activate the cube animation.",
                     "Let's get the cube animation underway, alright?",
                     "Okay, let's get the cube animation rolling."]
    response = random.choice(response_list)
    data_dic["play_animation"] = "play"
    print(data_dic)
    return response, data_dic

#Make text bold
def task8(most_similar_task):
    reset_all()
    response_list = ["I made that part bold.",
                     "I emphasized that part by making it bold.",
                     "I formatted that section to be bold."]
    response = random.choice(response_list)
    doc = BOLD_MODEL(most_similar_task)
    print(f"Text: {most_similar_task}")
    updated_task = most_similar_task
    for ent in doc.ents:
        print(f"{ent.text} ({ent.label_})")
        updated_task = updated_task.replace(ent.text, '')
    updated_task = updated_task.replace(".", '')
    print(f"Updated Text: {updated_task}")
    data_dic["bold_text"] = updated_task
    return response, data_dic

def task9(most_similar_task):
    reset_all()
    response_list = ["I wrote about it in a word file.",
                     "I documented it in a Word file.",
                     "I noted it down in a Word file."]
    response = random.choice(response_list)
    doc = TOPIC_MODEL(most_similar_task)
    print(f"Text: {most_similar_task}")
    updated_task = most_similar_task
    for ent in doc.ents:
        print(f"{ent.text} ({ent.label_})")
        updated_task = updated_task.replace(ent.text, '')
    updated_task = updated_task.replace(".", '').replace("?", '').replace("!", '')
    print(f"Updated Text: {updated_task}")
    prompt = "Please tell me about" + updated_task
    input_prompt = f"""[INST]
    <<SYS>> You are smart AI assistant
    <<SYS>> 
    {prompt} [INST] Please keep your answers short and simple: """
    output = llm("<user>\n" + input_prompt + "\n<assistant>\n", max_tokens=200)
    output = output['choices'][0]["text"]
    # Find the position of the last punctuation mark in a string
    last_punctuation_index = max(output.rfind('.'), output.rfind('?'), output.rfind('!'))
    # If a punctuation mark is found, remove subsequent characters and return
    if last_punctuation_index != -1:
        output = output[:last_punctuation_index + 1]
    # Create a new Document object
    document = Document()
    # Add title
    document.add_heading(updated_task, level=1)
    # Add paragraph
    document.add_paragraph(output)
    # Save file
    document.save('static/sample.docx')
    data_dic["word_file"] = updated_task
    return response, data_dic

# Dictionary of tasks and their corresponding functions
task_func_dict = {
    "Change the color of the [] button to [].": task1,
    "Please change the color of the [] button to [].": task1,
    "Please enter [] into the form.": task2,
    "I want to enter [] into the form.": task2,
    "Please change the position of the input form and yellow button.": task3,
    "Add a new button element.": task4,
    "I would like to add a new button element.": task4,
    "Delete a button element.": task5,
    "Delete a new button element.": task5,
    "Please play line animation.": task6,
    "play line animation.": task6,
    "Please stop animation.": task6_2,
    "stop animation.": task6_2,
    "Please play cube animation.": task7,
    "play cube animation.": task7,
    "Make the part [information] in bold.": task8,
    "Could you please write an article on the following [information]?": task9,
    "Please make word file about [information]?": task9,
}

@app.route("/main_agent")
def main_agent(message):
    task_to_compare = message
    # Calculate the degree of similarity between each task and the task it is compared to
    similarities = [compute_similarity(task, task_to_compare) for task in task_func_dict.keys()]
    # Extract tasks with similarity greater than 50%
    similar_tasks = [task for task, similarity in zip(task_func_dict.keys(), similarities) if similarity > 0.7]
    # Show most similar task
    if similar_tasks:
        most_similar_task = max(similar_tasks, key=lambda task: compute_similarity(task, task_to_compare))
        print(f"Most similar task: {most_similar_task}")
        # Execute the function of tasks whose similarity is above a certain level
        response, task_dic = task_func_dict[most_similar_task](task_to_compare)
        return response, task_dic
    else:
        prompt = task_to_compare
        input_prompt = f"""[INST]
        <<SYS>> You are smart AI assistant
        <<SYS>> 
        {prompt} [INST] Please keep your answers short and simple: """
        output = llm("<user>\n" + input_prompt + "\n<assistant>\n", max_tokens=200)
        output = output['choices'][0]["text"]
        # Find the position of the last punctuation mark in a string
        last_punctuation_index = max(output.rfind('.'), output.rfind('?'), output.rfind('!'))
        # If a punctuation mark is found, remove subsequent characters and return
        if last_punctuation_index != -1:
            output = output[:last_punctuation_index + 1]
        print(output)
        return output, None




